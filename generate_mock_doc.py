from docx import Document

def create_mock_document(filename):
    doc = Document()
    doc.add_heading('Ticket Log - Issue Tracker', 0)

    doc.add_heading('Ticket #101', level=1)
    doc.add_paragraph('Reported by: Rashi Patil (rashhi.patil@gmail.com)')
    doc.add_paragraph('Phone: +91 9876543210')
    doc.add_paragraph('DOB: 12/05/1990')
    doc.add_paragraph('Company: Acme Corp')
    doc.add_paragraph('Address: 123 Main St, Springfield, IL 62701')
    doc.add_paragraph('Issue: The server at 192.168.1.1 is not responding. Please check.')

    doc.add_heading('Ticket #102', level=1)
    doc.add_paragraph('Reported by: Rohan Dey (rohan.dey@gmail.com)')
    doc.add_paragraph('Phone: +91 9123456789')
    doc.add_paragraph('DOB: 1985-08-20')
    doc.add_paragraph('Company: Globex Corporation')
    doc.add_paragraph('Address: 456 Elm St, Shelbyville, IL 62702')
    doc.add_paragraph('Issue: Payment failed with credit card 4111 1111 1111 1111. User SSN is 000-12-3456. Please investigate.')

    doc.add_heading('Ticket #103', level=1)
    doc.add_paragraph('Reported by: John Doe (john.doe@example.com)')
    doc.add_paragraph('Phone: 555-1234')
    doc.add_paragraph('DOB: 01/01/2000')
    doc.add_paragraph('Company: Initech')
    doc.add_paragraph('Address: 789 Oak St, Capital City, IL 62703')
    doc.add_paragraph('Issue: Login failed from IP 10.0.0.1. Error code 500.')

    doc.save(filename)
    print(f"Mock document '{filename}' created successfully.")

if __name__ == '__main__':
    create_mock_document('mock_ticket_log.docx')
