import re
import argparse
from docx import Document
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig
from faker import Faker

# Setup Faker with a seed for reproducibility
faker = Faker()
Faker.seed(42)

# Global dictionary to map real entities to fake entities consistently
# e.g., "Rashi Patil" -> "John Doe"
ENTITY_MAP = {}

def get_fake_value(entity_type, real_value):
    """Generate a fake value based on the entity type, consistently mapping the same real value to the same fake value."""
    if real_value in ENTITY_MAP:
        return ENTITY_MAP[real_value]
        
    if entity_type == "PERSON":
        fake_val = faker.name()
    elif entity_type == "EMAIL_ADDRESS":
        fake_val = faker.email()
    elif entity_type == "PHONE_NUMBER":
        fake_val = faker.phone_number()
    elif entity_type == "LOCATION":
        fake_val = faker.address().replace('\n', ', ')
    elif entity_type == "DATE_TIME":
        fake_val = faker.date()
    elif entity_type == "US_SSN":
        fake_val = faker.ssn()
    elif entity_type == "CREDIT_CARD":
        fake_val = faker.credit_card_number()
    elif entity_type == "IP_ADDRESS":
        fake_val = faker.ipv4()
    elif entity_type == "URL":
        fake_val = faker.url()
    elif entity_type == "ORGANIZATION":
        fake_val = faker.company()
    else:
        fake_val = f"[{entity_type}]"
        
    ENTITY_MAP[real_value] = fake_val
    return fake_val


def redact_text(text, analyzer, anonymizer):
    """Analyze and anonymize text."""
    if not text.strip():
        return text

    # Analyze the text
    results = analyzer.analyze(text=text, entities=[], language='en')
    
    if not results:
        return text
        
    # We want to use our custom get_fake_value function. 
    # Presidio's AnonymizerEngine uses operators. We can create a custom operator logic by doing it manually
    # or by providing a custom operator. For simplicity and consistent mapping, we'll replace manually from the end to start.
    
    # Sort results by start position in descending order to avoid messing up indices
    results.sort(key=lambda x: x.start, reverse=True)
    
    redacted_text = text
    for result in results:
        real_value = text[result.start:result.end]
        fake_value = get_fake_value(result.entity_type, real_value)
        redacted_text = redacted_text[:result.start] + fake_value + redacted_text[result.end:]
        
    return redacted_text

def redact_docx(input_path, output_path):
    """Read a docx file, redact PII, and save to a new file."""
    # Initialize Presidio
    analyzer = AnalyzerEngine()
    anonymizer = AnonymizerEngine()
    
    doc = Document(input_path)
    
    # Redact paragraphs
    for para in doc.paragraphs:
        if para.text:
            redacted = redact_text(para.text, analyzer, anonymizer)
            # A simple replacement for the whole paragraph text.
            # This might lose some inline formatting (bold/italic on specific words)
            # but is the most robust way to replace text across runs.
            if redacted != para.text:
                para.text = redacted
                
    # Redact tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        redacted = redact_text(para.text, analyzer, anonymizer)
                        if redacted != para.text:
                            para.text = redacted
                            
    doc.save(output_path)
    print(f"Redaction complete. Saved to {output_path}")
    print("\nEntity Mapping Used:")
    for real, fake in ENTITY_MAP.items():
        print(f"  {real} -> {fake}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Redact PII from a docx file.")
    parser.add_argument("input_file", help="Path to the input docx file")
    parser.add_argument("output_file", help="Path to the output docx file")
    args = parser.parse_args()
    
    redact_docx(args.input_file, args.output_file)
